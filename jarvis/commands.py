"""
Módulo de comandos del sistema: abrir apps, buscar en internet, controlar volumen, apagar PC.
"""

import subprocess
import webbrowser
import platform
import os
import shutil
import re
import threading
import time
from datetime import datetime
from pathlib import Path
from urllib.request import Request, urlopen
from urllib.parse import quote, quote_plus

from jarvis.interaction_memory import (
    build_media_confirmation,
    build_search_confirmation,
    build_site_confirmation,
)
from jarvis.gmail_service import read_inbox, send_email
from jarvis.settings import get_setting


def _default_documents_dir() -> Path:
    home = Path.home()
    for folder_name in ("Documents", "Documentos"):
        candidate = home / folder_name
        if candidate.exists():
            return candidate
    return home


def _safe_filename(name: str, fallback: str = "documento") -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._ -]+", "", (name or "").strip())
    cleaned = cleaned.replace(" ", "_").strip("._")
    return cleaned or fallback


def _is_heading_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith(("# ", "## ", "### ")):
        return True
    return stripped.endswith(":") and len(stripped) <= 90


def _heading_level(line: str) -> int:
    stripped = line.strip()
    if stripped.startswith("### "):
        return 3
    if stripped.startswith("## "):
        return 2
    return 1


def _clean_heading_text(line: str) -> str:
    stripped = line.strip()
    if stripped.startswith("### "):
        return stripped[4:].strip()
    if stripped.startswith("## "):
        return stripped[3:].strip()
    if stripped.startswith("# "):
        return stripped[2:].strip()
    return stripped.rstrip(":").strip()


def _write_structured_content(document, content: str) -> None:
    pending_paragraph: list[str] = []

    def flush_paragraph() -> None:
        if not pending_paragraph:
            return
        paragraph = document.add_paragraph(" ".join(part.strip() for part in pending_paragraph if part.strip()))
        paragraph.style = "Normal"
        pending_paragraph.clear()

    for raw_line in content.splitlines():
        line = raw_line.strip()

        if not line:
            flush_paragraph()
            continue

        if _is_heading_line(line):
            flush_paragraph()
            document.add_heading(_clean_heading_text(line), level=_heading_level(line))
            continue

        if line.startswith(("- ", "* ", "• ")):
            flush_paragraph()
            bullet = document.add_paragraph(style="List Bullet")
            bullet.add_run(line[2:].strip())
            continue

        if re.match(r"^\d+[\.)]\s+", line):
            flush_paragraph()
            numbered = document.add_paragraph(style="List Number")
            numbered.add_run(re.sub(r"^\d+[\.)]\s+", "", line).strip())
            continue

        pending_paragraph.append(line)

    flush_paragraph()


def create_word_document(title: str, content: str) -> str:
    """Crea un archivo .docx con el contenido solicitado y lo abre."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        return "Falta instalar python-docx. Ejecuta: pip install python-docx"

    title = (title or "Documento JARVIS").strip()
    content = (content or "").strip()
    if not content:
        return "No recibi contenido para crear el documento de Word."

    docs_dir = _default_documents_dir()
    docs_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"{_safe_filename(title, fallback='documento_jarvis')}_{timestamp}.docx"
    doc_path = docs_dir / file_name
    author_name = str(get_setting("user_name") or "Usuario").strip() or "Usuario"
    human_date = datetime.now().strftime("%d/%m/%Y %H:%M")

    document = Document()
    document.core_properties.author = author_name
    document.core_properties.title = title

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(22)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(f"Documento generado por {author_name}")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(human_date)
    date_run.font.size = Pt(10)

    document.add_paragraph("")
    _write_structured_content(document, content)
    document.save(doc_path)

    try:
        system = platform.system()
        if system == "Windows":
            os.startfile(str(doc_path))
        elif system == "Darwin":
            subprocess.Popen(["open", str(doc_path)])
        else:
            opener = shutil.which("xdg-open")
            if opener:
                subprocess.Popen([opener, str(doc_path)])
    except Exception:
        pass

    return f"Listo. Cree el documento de Word '{title}' y lo guarde en {doc_path}."


def create_word_document_from_payload(payload: str) -> str:
    """Procesa DOC_CREATE usando primera linea como titulo y el resto como contenido."""
    raw = (payload or "").strip()
    if not raw:
        return "No recibi datos para crear el documento."

    if "\n" in raw:
        title_line, content = raw.split("\n", 1)
        title = title_line.strip() or "Documento JARVIS"
        return create_word_document(title, content)

    if "|" in raw:
        title, content = raw.split("|", 1)
        return create_word_document(title, content)

    return create_word_document("Documento JARVIS", raw)


def parse_media_request(user_text: str) -> str | None:
    """Detecta peticiones musicales comunes en lenguaje natural y devuelve payload MEDIA_PLAY."""
    raw = (user_text or "").strip()
    if not raw:
        return None

    lowered = raw.lower()
    media_triggers = (
        "pon ",
        "ponme ",
        "reproduce ",
        "reproducir ",
        "quiero escuchar",
        "escucha ",
        "busca musica",
        "busca la cancion",
        "pon la cancion",
        "pon musica",
        "pon un album",
        "abre spotify y pon",
        "abre youtube y pon",
        "abre youtube music y pon",
    )
    if not any(trigger in lowered for trigger in media_triggers):
        return None

    platform_name = "youtube"
    platform_patterns = (
        ("youtube music", "youtube_music"),
        ("youtube music", "youtube_music"),
        ("yt music", "youtube_music"),
        ("spotify", "spotify"),
        ("youtube", "youtube"),
        ("yt", "youtube"),
    )
    for marker, mapped in platform_patterns:
        if marker in lowered:
            platform_name = mapped
            lowered = lowered.replace(marker, " ")
            raw = re.sub(marker, " ", raw, flags=re.IGNORECASE)
            break

    query = raw
    cleanup_patterns = [
        r"\babre\b",
        r"\bspotify\b",
        r"\byoutube\s+music\b",
        r"\byt\s+music\b",
        r"\byoutube\b",
        r"\byt\b",
        r"\bponme\b",
        r"\bpon\b",
        r"\breproduce\b",
        r"\breproducir\b",
        r"\bquiero escuchar\b",
        r"\bescuchar\b",
        r"\bescucha\b",
        r"\bbusca\b",
        r"\bmusica\b",
        r"\bmúsica\b",
        r"\bla cancion\b",
        r"\bla canción\b",
        r"\bcancion\b",
        r"\bcanción\b",
        r"\bdel artista\b",
        r"\bdel grupo\b",
        r"\bde la banda\b",
        r"\ben spotify\b",
        r"\ben youtube music\b",
        r"\ben youtube\b",
    ]
    for pattern in cleanup_patterns:
        query = re.sub(pattern, " ", query, flags=re.IGNORECASE)

    query = re.sub(r"\s+", " ", query).strip(" ,.;:-")
    query = re.sub(r"^(y|e)\s+", "", query, flags=re.IGNORECASE)
    query = re.sub(r"\s+(y|e|en)$", "", query, flags=re.IGNORECASE)
    query = re.sub(r"\s+", " ", query).strip(" ,.;:-")
    if len(query) < 2:
        return None

    return f"{platform_name}|{query}"


def _find_youtube_first_video_url(query: str) -> str | None:
    """Busca el primer video en resultados de YouTube y devuelve su URL."""
    try:
        search_url = f"https://www.youtube.com/results?search_query={quote_plus(query)}"
        request = Request(
            search_url,
            headers={
                "User-Agent": "Mozilla/5.0",
                "Accept-Language": "es-ES,es;q=0.9,en;q=0.8",
            },
        )
        with urlopen(request, timeout=8) as response:
            html = response.read().decode("utf-8", errors="ignore")
    except Exception:
        return None

    match = re.search(r'"videoId":"([^"]{11})"', html)
    if not match:
        return None
    return f"https://www.youtube.com/watch?v={match.group(1)}"


def _run_desktop_automation(task) -> bool:
    """Ejecuta automatizacion de escritorio de forma segura y no bloqueante."""
    try:
        import pyautogui  # noqa: F401
    except Exception:
        return False

    threading.Thread(target=task, daemon=True).start()
    return True


def _spotify_autoplay_task(query: str) -> None:
    try:
        import pyautogui

        time.sleep(2.8)
        pyautogui.hotkey("ctrl", "l")
        time.sleep(0.3)
        pyautogui.write(query, interval=0.02)
        pyautogui.press("enter")
        time.sleep(1.4)
        pyautogui.press("tab", presses=3, interval=0.15)
        pyautogui.press("enter")
    except Exception:
        return


def _youtube_music_autoplay_task() -> None:
    try:
        import pyautogui

        time.sleep(3.0)
        pyautogui.press("tab", presses=6, interval=0.12)
        pyautogui.press("enter")
    except Exception:
        return


def open_application(app_name: str) -> str:
    """Intenta abrir una aplicación por nombre."""
    app_name = app_name.strip().lower()
    system = platform.system()

    # Mapa de nombres comunes → ejecutables
    app_map = {
        # Navegadores
        "chrome": "chrome.exe",
        "google chrome": "chrome.exe",
        "edge": "msedge.exe",
        "microsoft edge": "msedge.exe",
        "firefox": "firefox.exe",
        "brave": "brave.exe",
        "opera": "opera.exe",
        # Música / entretenimiento
        "spotify": "spotify.exe",
        "vlc": "vlc.exe",
        # Comunicación
        "discord": "discord.exe",
        "whatsapp": "whatsapp.exe",
        "telegram": "telegram.exe",
        "zoom": "zoom.exe",
        "teams": "teams.exe",
        "microsoft teams": "teams.exe",
        "skype": "skype.exe",
        # Ofimática
        "word": "winword.exe",
        "excel": "excel.exe",
        "powerpoint": "powerpnt.exe",
        "outlook": "outlook.exe",
        # Herramientas del sistema
        "notepad": "notepad.exe",
        "bloc de notas": "notepad.exe",
        "calculadora": "calc.exe",
        "calculator": "calc.exe",
        "explorador": "explorer.exe",
        "explorador de archivos": "explorer.exe",
        "explorer": "explorer.exe",
        "paint": "mspaint.exe",
        "cmd": "cmd.exe",
        "terminal": "cmd.exe",
        "administrador de tareas": "taskmgr.exe",
        "task manager": "taskmgr.exe",
        "panel de control": "control.exe",
        "configuracion": "ms-settings:",
        "configuración": "ms-settings:",
        # Desarrollo
        "vscode": "code",
        "visual studio code": "code",
        # Juegos
        "steam": "steam.exe",
        "epic games": "epicgameslauncher.exe",
    }

    linux_map = {
        # Navegadores
        "chrome": "google-chrome",
        "google chrome": "google-chrome",
        "edge": "microsoft-edge",
        "microsoft edge": "microsoft-edge",
        "firefox": "firefox",
        "brave": "brave-browser",
        "opera": "opera",
        # Musica / entretenimiento
        "spotify": "spotify",
        "vlc": "vlc",
        # Comunicacion
        "discord": "discord",
        "whatsapp": "https://web.whatsapp.com",
        "telegram": "telegram-desktop",
        "zoom": "zoom",
        "teams": "teams-for-linux",
        "microsoft teams": "teams-for-linux",
        "skype": "skypeforlinux",
        # Herramientas del sistema
        "terminal": "x-terminal-emulator",
        "explorador": "xdg-open .",
        "explorador de archivos": "xdg-open .",
        "explorer": "xdg-open .",
        # Desarrollo
        "vscode": "code",
        "visual studio code": "code",
        # Juegos
        "steam": "steam",
    }

    if system == "Linux":
        executable = linux_map.get(app_name, app_name)
    else:
        executable = app_map.get(app_name, app_name)

    try:
        if system == "Windows":
            # En Windows, varios navegadores no siempre están en PATH.
            # Usamos protocolos/START para que el sistema resuelva la app instalada.
            if executable == "msedge.exe":
                os.startfile("microsoft-edge:")
            elif executable in ("chrome.exe", "firefox.exe", "brave.exe", "opera.exe"):
                subprocess.Popen(["cmd", "/c", "start", "", executable], shell=True)
            elif executable == "ms-settings:":
                os.startfile("ms-settings:")
            else:
                subprocess.Popen(executable, shell=True)
        elif system == "Darwin":  # macOS
            subprocess.Popen(["open", "-a", executable])
        else:  # Linux
            if executable.startswith("http://") or executable.startswith("https://"):
                webbrowser.open(executable)
            elif executable.startswith("xdg-open "):
                subprocess.Popen(executable, shell=True)
            else:
                if shutil.which(executable) is None:
                    return f"No encontre la app '{app_name}' en Linux (comando: {executable})."
                subprocess.Popen([executable])
        return f"Abriendo {app_name}..."
    except Exception as e:
        return f"No pude abrir {app_name}. Error: {e}"


def open_multiple_applications(payload: str) -> str:
    """Abre varias apps en una sola orden. Formato sugerido: app1|app2|app3."""
    raw = (payload or "").strip()
    if not raw:
        return "No recibi aplicaciones para abrir."

    if "|" in raw:
        parts = [p.strip() for p in raw.split("|") if p.strip()]
    else:
        parts = [p.strip() for p in re.split(r"\s*(?:,| y | e | and )\s*", raw, flags=re.IGNORECASE) if p.strip()]

    if not parts:
        return "No pude identificar aplicaciones para abrir."

    opened: list[str] = []
    failed: list[str] = []

    for app in parts:
        result = open_application(app)
        if result.lower().startswith("abriendo "):
            opened.append(app)
        else:
            failed.append(f"{app}: {result}")

    if opened and not failed:
        return f"Listo, abri estas apps: {', '.join(opened)}."
    if opened and failed:
        return (
            f"Abri: {', '.join(opened)}. "
            f"No pude abrir: {' | '.join(failed)}"
        )
    return f"No pude abrir ninguna app. Detalle: {' | '.join(failed)}"


def web_search(query: str) -> str:
    """Abre una búsqueda en el navegador predeterminado."""
    query = query.strip()
    url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
    webbrowser.open(url)
    return build_search_confirmation(query)


def open_website(target: str) -> str:
    """Abre un sitio web directo en el navegador predeterminado."""
    target = target.strip().lower()

    site_map = {
        "instagram": "https://www.instagram.com",
        "facebook": "https://www.facebook.com",
        "whatsapp": "https://web.whatsapp.com",
        "whatsapp web": "https://web.whatsapp.com",
        "gmail": "https://mail.google.com",
        "youtube": "https://www.youtube.com",
        "x": "https://x.com",
        "twitter": "https://x.com",
        "tiktok": "https://www.tiktok.com",
        "linkedin": "https://www.linkedin.com",
        "github": "https://github.com",
        "netflix": "https://www.netflix.com",
    }

    url = site_map.get(target, target)
    if not url.startswith(("http://", "https://")):
        url = f"https://{url}"

    webbrowser.open(url)
    return build_site_confirmation(target)


def play_media(payload: str) -> str:
    """Abre una reproduccion o busqueda multimedia en YouTube, YouTube Music o Spotify."""
    raw = (payload or "").strip()
    platform_name = "youtube"
    query = raw

    if "|" in raw:
        maybe_platform, maybe_query = [part.strip() for part in raw.split("|", 1)]
        if maybe_query:
            platform_name = maybe_platform.lower() or "youtube"
            query = maybe_query

    query = query.strip()
    platform_name = platform_name.strip().lower() or "youtube"

    if not query:
        return "No recibi el nombre de la cancion, artista o video para reproducir."

    platform_aliases = {
        "yt": "youtube",
        "youtube.com": "youtube",
        "youtube music": "youtube_music",
        "yt music": "youtube_music",
        "music youtube": "youtube_music",
        "youtube_music": "youtube_music",
        "spotify web": "spotify",
    }
    platform_name = platform_aliases.get(platform_name, platform_name)

    if platform_name == "spotify":
        spotify_uri = f"spotify:search:{query}"
        spotify_web = f"https://open.spotify.com/search/{quote(query)}"
        opened = webbrowser.open(spotify_uri)
        if not opened:
            webbrowser.open(spotify_web)
        automated = _run_desktop_automation(lambda: _spotify_autoplay_task(query))
        if automated:
            return f"Intente reproducir {query} en Spotify automaticamente. Si no arranca solo, deja Spotify abierto y repitemelo." 
        return build_media_confirmation("spotify", query)

    if platform_name == "youtube":
        youtube_url = _find_youtube_first_video_url(query) or f"https://www.youtube.com/results?search_query={quote_plus(query)}"
        webbrowser.open(youtube_url)
        if "watch?v=" in youtube_url:
            return f"Reproduciendo {query} en YouTube." 
        return build_media_confirmation("youtube", query)

    if platform_name == "youtube_music":
        youtube_url = f"https://music.youtube.com/search?q={quote_plus(query)}"
        webbrowser.open(youtube_url)
        automated = _run_desktop_automation(_youtube_music_autoplay_task)
        if automated:
            return f"Intente reproducir {query} en YouTube Music automaticamente." 
        return build_media_confirmation("youtube music", query)

    generic_url = f"https://www.google.com/search?q={quote_plus(query + ' ' + platform_name)}"
    webbrowser.open(generic_url)
    return f"No reconoci la plataforma {platform_name}. Te deje una busqueda web lista para {query}."


def gmail_inbox(arg: str) -> str:
    """Revisa la bandeja de Gmail y da un resumen."""
    raw = (arg or "").strip()
    max_results = 5
    if raw:
        try:
            max_results = int(raw)
        except ValueError:
            max_results = 5

    # Tambien abre Gmail para que el usuario vea la bandeja al instante.
    webbrowser.open("https://mail.google.com/mail/u/0/#inbox")
    return read_inbox(max_results=max_results)


def gmail_send(payload: str) -> str:
    """Envia correo. Formato: to|subject|body"""
    parts = [p.strip() for p in (payload or "").split("|", 2)]
    if len(parts) < 3:
        return (
            "Formato de envio invalido. Usa: MAIL_SEND:destinatario@correo.com|Asunto|Mensaje"
        )

    to_email, subject, body = parts
    return send_email(to_email=to_email, subject=subject, body=body)


def control_volume(action: str) -> str:
    """Controla el volumen del sistema (Windows/Linux)."""
    action = action.strip().lower()
    system = platform.system()

    if action not in {"up", "down", "mute"}:
        return f"Accion de volumen no reconocida: {action}"

    if system == "Linux":
        try:
            if shutil.which("wpctl"):
                if action == "up":
                    subprocess.Popen(["wpctl", "set-volume", "@DEFAULT_AUDIO_SINK@", "5%+"])
                    return "Volumen subido (Linux/wpctl)."
                if action == "down":
                    subprocess.Popen(["wpctl", "set-volume", "@DEFAULT_AUDIO_SINK@", "5%-"])
                    return "Volumen bajado (Linux/wpctl)."
                subprocess.Popen(["wpctl", "set-mute", "@DEFAULT_AUDIO_SINK@", "toggle"])
                return "Mute alternado (Linux/wpctl)."

            if shutil.which("pactl"):
                if action == "up":
                    subprocess.Popen(["pactl", "set-sink-volume", "@DEFAULT_SINK@", "+5%"])
                    return "Volumen subido (Linux/pactl)."
                if action == "down":
                    subprocess.Popen(["pactl", "set-sink-volume", "@DEFAULT_SINK@", "-5%"])
                    return "Volumen bajado (Linux/pactl)."
                subprocess.Popen(["pactl", "set-sink-mute", "@DEFAULT_SINK@", "toggle"])
                return "Mute alternado (Linux/pactl)."

            if shutil.which("amixer"):
                if action == "up":
                    subprocess.Popen(["amixer", "-D", "pulse", "sset", "Master", "5%+"])
                    return "Volumen subido (Linux/amixer)."
                if action == "down":
                    subprocess.Popen(["amixer", "-D", "pulse", "sset", "Master", "5%-"])
                    return "Volumen bajado (Linux/amixer)."
                subprocess.Popen(["amixer", "-D", "pulse", "sset", "Master", "toggle"])
                return "Mute alternado (Linux/amixer)."

            return "No encontre utilidades de audio (wpctl/pactl/amixer) en Linux."
        except Exception as e:
            return f"Error controlando volumen en Linux: {e}"

    if system != "Windows":
        return "Control de volumen disponible en Windows y Linux."

    try:
        from ctypes import cast, POINTER
        from comtypes import CLSCTX_ALL
        from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume

        devices = AudioUtilities.GetSpeakers()
        interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
        volume = cast(interface, POINTER(IAudioEndpointVolume))

        current = volume.GetMasterVolumeLevelScalar()

        if action == "up":
            new_vol = min(1.0, current + 0.1)
            volume.SetMasterVolumeLevelScalar(new_vol, None)
            return f"Volumen subido al {int(new_vol * 100)}%"
        elif action == "down":
            new_vol = max(0.0, current - 0.1)
            volume.SetMasterVolumeLevelScalar(new_vol, None)
            return f"Volumen bajado al {int(new_vol * 100)}%"
        elif action == "mute":
            muted = volume.GetMute()
            volume.SetMute(not muted, None)
            return "Silenciado" if not muted else "Sonido activado"
    except ImportError:
        return "Instala 'pycaw' para controlar el volumen: pip install pycaw"
    except Exception as e:
        return f"Error controlando volumen: {e}"


def system_power(action: str) -> str:
    """Controla el estado del sistema: apagar, reiniciar, suspender, bloquear."""
    action = action.strip().lower()
    system = platform.system()

    if system == "Linux":
        try:
            if action == "shutdown":
                if shutil.which("shutdown"):
                    subprocess.Popen(["shutdown", "-h", "+1", "ECHONEX: apagado solicitado"])
                    return "Apagando Linux en 1 minuto. Puedes cancelar con POWER:cancel_shutdown."
                if shutil.which("systemctl"):
                    subprocess.Popen(["systemctl", "poweroff"])
                    return "Apagando Linux ahora (systemctl)."
                return "No encontre comando para apagar en Linux."

            if action == "restart":
                if shutil.which("shutdown"):
                    subprocess.Popen(["shutdown", "-r", "+1", "ECHONEX: reinicio solicitado"])
                    return "Reiniciando Linux en 1 minuto. Puedes cancelar con POWER:cancel_shutdown."
                if shutil.which("systemctl"):
                    subprocess.Popen(["systemctl", "reboot"])
                    return "Reiniciando Linux ahora (systemctl)."
                return "No encontre comando para reiniciar en Linux."

            if action == "sleep":
                if shutil.which("systemctl"):
                    subprocess.Popen(["systemctl", "suspend"])
                    return "Suspendiendo Linux..."
                return "No encontre systemctl para suspender en Linux."

            if action == "lock":
                if shutil.which("loginctl"):
                    subprocess.Popen(["loginctl", "lock-session"])
                    return "Bloqueando la pantalla en Linux..."
                if shutil.which("gnome-screensaver-command"):
                    subprocess.Popen(["gnome-screensaver-command", "-l"])
                    return "Bloqueando la pantalla en Linux..."
                if shutil.which("xdg-screensaver"):
                    subprocess.Popen(["xdg-screensaver", "lock"])
                    return "Bloqueando la pantalla en Linux..."
                return "No encontre comando para bloquear pantalla en Linux."

            if action == "unlock":
                return "Linux no permite desbloquear automaticamente sin credenciales. Debes desbloquear manualmente."

            if action == "cancel_shutdown":
                if shutil.which("shutdown"):
                    subprocess.Popen(["shutdown", "-c"])
                    return "Apagado/reinicio cancelado en Linux."
                return "No encontre comando shutdown para cancelar en Linux."

            return f"Accion de energia no reconocida: {action}"
        except Exception as e:
            return f"Error ejecutando accion de energia en Linux: {e}"

    if system != "Windows":
        return "Control de energia disponible en Windows y Linux."

    if action == "shutdown":
        subprocess.Popen("shutdown /s /t 10 /c \"ECHONEX: Apagando el equipo...\"", shell=True)
        return "Apagando el equipo en 10 segundos. Escribe 'shutdown /a' en CMD para cancelar."
    elif action == "restart":
        subprocess.Popen("shutdown /r /t 10 /c \"ECHONEX: Reiniciando el equipo...\"", shell=True)
        return "Reiniciando el equipo en 10 segundos. Escribe 'shutdown /a' en CMD para cancelar."
    elif action == "sleep":
        subprocess.Popen("rundll32.exe powrprof.dll,SetSuspendState 0,1,0", shell=True)
        return "Suspendiendo el equipo..."
    elif action == "lock":
        subprocess.Popen("rundll32.exe user32.dll,LockWorkStation", shell=True)
        return "Bloqueando la pantalla..."
    elif action == "unlock":
        return (
            "Windows no permite desbloquear automaticamente sin credenciales (PIN/huella/password). "
            "Por seguridad, debes desbloquear manualmente."
        )
    elif action == "cancel_shutdown":
        subprocess.Popen("shutdown /a", shell=True)
        return "Apagado/reinicio cancelado."
    return f"Acción de energía no reconocida: {action}"


def _execute_single_command(command_text: str) -> str | None:
    command_text = command_text.strip()
    if not command_text:
        return None

    if command_text.startswith("OPEN_APPS:"):
        payload = command_text.split("OPEN_APPS:", 1)[1].strip()
        return open_multiple_applications(payload)

    if command_text.startswith("OPEN_APP:"):
        app = command_text.split("OPEN_APP:", 1)[1].strip()
        if "|" in app or re.search(r"\s(?:y|e|and)\s", app, flags=re.IGNORECASE):
            return open_multiple_applications(app)
        return open_application(app)

    if command_text.startswith("WEB_SEARCH:"):
        query = command_text.split("WEB_SEARCH:", 1)[1].strip()
        return web_search(query)

    if command_text.startswith("WEB_OPEN:"):
        target = command_text.split("WEB_OPEN:", 1)[1].strip()
        return open_website(target)

    if command_text.startswith("MEDIA_PLAY:"):
        payload = command_text.split("MEDIA_PLAY:", 1)[1].strip()
        return play_media(payload)

    if command_text.startswith("VOLUME:"):
        action = command_text.split("VOLUME:", 1)[1].strip()
        return control_volume(action)

    if command_text.startswith("POWER:"):
        action = command_text.split("POWER:", 1)[1].strip()
        return system_power(action)

    if command_text.startswith("MAIL_INBOX:"):
        arg = command_text.split("MAIL_INBOX:", 1)[1].strip()
        return gmail_inbox(arg)

    if command_text.startswith("MAIL_SEND:"):
        payload = command_text.split("MAIL_SEND:", 1)[1].strip()
        return gmail_send(payload)

    if command_text.startswith("DOC_CREATE:"):
        payload = command_text.split("DOC_CREATE:", 1)[1]
        return create_word_document_from_payload(payload)

    return None


def handle_command(response: str) -> str | None:
    """
    Detecta si la respuesta del LLM es un comando especial y lo ejecuta.
    Retorna la respuesta de texto para hablar, o None si no hubo comando.
    """
    response = response.strip()

    if response.startswith("ACTIONS:"):
        payload = response.split("ACTIONS:", 1)[1].strip()
        actions = [a.strip() for a in re.split(r"\s*;;\s*", payload) if a.strip()]
        if not actions:
            return "No recibi acciones para ejecutar."

        summaries: list[str] = []
        for idx, action in enumerate(actions, start=1):
            result = _execute_single_command(action)
            if result:
                summaries.append(f"Accion {idx}: {result}")
            else:
                summaries.append(f"Accion {idx}: no reconocida ({action}).")

        return " ".join(summaries)

    return _execute_single_command(response)
